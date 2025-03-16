from dataclasses import dataclass
from datetime import datetime
import zipfile
import logging
from typing import Dict, List, Optional
from docx.api import Document
from xml.etree import ElementTree
from setting import (
    ESSAY_PROMPT_KEY,
    ESSAY_TEXT_KEY,
    COMMENTS_KEY,
    GENERAL_FEEDBACK_KEY,
    COMMENT_ID_KEY,
    COMMENT_START_KEY,
    COMMENT_END_KEY,
    COMMENT_TEXT_KEY,
    HIGHLIGHTED_TEXT_KEY,
    AUTHOR_KEY,
    DATE_KEY,
)

logger = logging.getLogger(__name__)


@dataclass
class Comment:
    id: str
    para_id: str
    para_id_parent: Optional[str]
    author: str
    date: str
    comment_text: str
    highlighted_text: str
    start: int
    end: int

    def get_dict(
        self, include_author: bool = False, include_date: bool = False
    ) -> Dict:
        d = {
            COMMENT_ID_KEY: self.id,
            COMMENT_START_KEY: self.start,
            COMMENT_END_KEY: self.end,
            HIGHLIGHTED_TEXT_KEY: self.highlighted_text,
            COMMENT_TEXT_KEY: self.comment_text,
        }
        if include_author:
            d[AUTHOR_KEY] = self.author
        if include_date:
            d[DATE_KEY] = self.date
        return d


class HighlightRange:
    """Highlight range with relative start position."""

    def __init__(self, comment_id: str, absolute_start: int):
        self.comment_id = comment_id
        self.absolute_start = absolute_start
        self.section_start = 0
        self.texts: List[str] = []

    def append(self, text: str):
        self.texts.append(text)

    def get_text(self) -> str:
        return "".join(self.texts)

    def get_relative_start(self) -> int:
        return self.absolute_start - self.section_start


@dataclass
class Section:
    start: int
    end: int
    raw_text: str
    stripped_text: str


class ExtractConfig:
    def __init__(
        self, 
        prompt_start_token=None,
        prompt_end_token=None, 
        fb_start_token=None, 
        fb_end_token=None, 
        include_author=False, 
        include_date=False
    ):
        self.prompt_start_token = prompt_start_token
        self.prompt_end_token = prompt_end_token
        self.fb_start_token = fb_start_token
        self.fb_end_token = fb_end_token
        self.include_author = include_author
        self.include_date = include_date


class CommentExtractor:
    def __init__(self, config: ExtractConfig = None):
        self.config = config or ExtractConfig()
        self.namespaces = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
            "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
        }
        self.new_line = "\n"

    def extract_text_between_tokens(self, text: str, start_token: str|None, end_token: str|None) -> Section:
        """Extract text between start and end tokens."""

        if start_token is None:
            start_idx = 0
        else:
            start_token_pos = text.find(start_token)
            if start_token_pos < 0:
                logger.warning(
                    "Start token not found in text: %s", start_token
                )
                start_idx = 0
            else:
                start_idx = start_token_pos + len(start_token)
        if end_token is None:
            end_idx = len(text)
        else:
            end_token_pos = text.find(end_token, start_idx)
            if end_token_pos < 0:
                logger.warning("End token not found in text: %s", end_token)
                end_idx = len(text)
            else:
                end_idx = end_token_pos

        raw_text = text[start_idx:end_idx]
        lstripped_text = raw_text.lstrip()
        stripped_text = lstripped_text.rstrip()
        blank_chars_before_start_token = len(raw_text) - len(lstripped_text)
        blank_chars_before_end_token = len(lstripped_text) - len(stripped_text)

        return Section(
            start=start_idx + blank_chars_before_start_token,
            end=end_idx - blank_chars_before_end_token,
            raw_text=raw_text,
            stripped_text=stripped_text,
        )

    def _read_docx_file(
        self, file_path: str
    ) -> tuple[ElementTree.Element, ElementTree.Element, ElementTree.Element]:
        try:
            with zipfile.ZipFile(file_path) as zip_ref:
                # Read XML files
                comments_root = self._read_xml_files(zip_ref, "word/comments.xml")
                comments_extend_root = self._read_xml_files(
                    zip_ref, "word/commentsExtended.xml"
                )
                doc_root = self._read_xml_files(zip_ref, "word/document.xml")
                return comments_root, comments_extend_root, doc_root
        except zipfile.BadZipFile as e:
            logger.error("Error reading %s: %s", file_path, str(e))
            return None, None, None

    def _read_xml_files(self, zip_ref, inner_file_name: str) -> ElementTree.Element:
        """Read and parse XML files from the Word document."""
        try:
            xml = zip_ref.read(inner_file_name)
        except KeyError:
            logger.warning(
                f"XML file {inner_file_name} not found in document {zip_ref.filename}"
            )
            return None
        return ElementTree.fromstring(xml)

    def _extract_highlight_ranges(self, doc_root) -> Dict[str, HighlightRange]:
        """Extract comment ranges and their corresponding text from document with position info."""
        current_range_ids = []
        comment_id_to_range = {}
        full_text = ""
        is_first_paragraph = True

        if doc_root is None:
            return comment_id_to_range, full_text

        for elem in doc_root.iter():
            if elem.tag == f"{{{self.namespaces['w']}}}commentRangeStart":
                comment_id = elem.get(f"{{{self.namespaces['w']}}}id")
                current_range_ids.append(comment_id)
                hr = HighlightRange(
                    comment_id=comment_id, absolute_start=len(full_text)
                )
                comment_id_to_range[comment_id] = hr
            elif elem.tag == f"{{{self.namespaces['w']}}}commentRangeEnd":
                comment_id = elem.get(f"{{{self.namespaces['w']}}}id")
                if comment_id in comment_id_to_range:
                    current_range_ids.remove(comment_id)
            elif elem.tag == f"{{{self.namespaces['w']}}}t":
                if elem.text:
                    for comment_id in current_range_ids:
                        hr = comment_id_to_range[comment_id]
                        hr.append(elem.text)
                    full_text += elem.text
            elif elem.tag == f"{{{self.namespaces['w']}}}p":
                if is_first_paragraph:
                    is_first_paragraph = False
                else:
                    # Add new line to all highlight ranges, if not the first paragraph
                    full_text += self.new_line
                    for comment_id in current_range_ids:
                        hr = comment_id_to_range[comment_id]
                        hr.append(self.new_line)
            elif elem.tag == f"{{{self.namespaces['w']}}}br":
                full_text += self.new_line
                for comment_id in current_range_ids:
                    hr = comment_id_to_range[comment_id]
                    hr.append(self.new_line)

        return comment_id_to_range, full_text

    def _get_sub_comments(self, comments_extend_root) -> set[str]:
        """Get set of paragraph IDs that are replies to other comments."""
        sub_comments = set()
        if comments_extend_root is None:
            return sub_comments
        for comment_ex in comments_extend_root.findall(
            ".//w15:commentEx", self.namespaces
        ):
            para_id = comment_ex.get(f"{{{self.namespaces['w15']}}}paraId")
            parent_para_id = comment_ex.get(
                f"{{{self.namespaces['w15']}}}paraIdParent", None
            )
            if parent_para_id is not None:
                sub_comments.add(para_id)
        return sub_comments

    def _extract_comment_text(self, comment_node) -> str:
        """Extract text from a comment node."""
        comment_text = []
        for p in comment_node.findall(".//w:t", self.namespaces):
            if p.text:
                comment_text.append(p.text)
        return " ".join(comment_text)

    def _extract_comments(
        self, comments_root, comment_id_to_range, section_start, sub_comments
    ) -> List[Dict]:
        # Process main comments
        comments = []

        for comment_node in comments_root.findall(".//w:comment", self.namespaces):
            comment_id = comment_node.get(f"{{{self.namespaces['w']}}}id")
            para = comment_node.find(".//w:p", self.namespaces)

            if para is None:
                continue

            para_id = para.get(f"{{{self.namespaces['w14']}}}paraId")
            if para_id in sub_comments:
                continue  # Skip reply comments

            highlighted_range = comment_id_to_range.get(comment_id, None)
            if highlighted_range is None:
                logger.warning(
                    "Skip comment. No highlighted text found for comment %s", comment_id
                )
                continue

            highlighted_range.section_start = section_start

            pos = highlighted_range.get_relative_start()
            if pos < 0:
                logger.warning(
                    "Skip comment. Comment %s appears before start token", comment_id
                )
                continue

            comment = Comment(
                id=int(comment_id),
                para_id=para_id,
                para_id_parent=None,
                author=comment_node.get(f"{{{self.namespaces['w']}}}author", ""),
                date=comment_node.get(
                    f"{{{self.namespaces['w']}}}date", datetime.now().isoformat()
                ),
                comment_text=self._extract_comment_text(comment_node),
                highlighted_text=highlighted_range.get_text(),
                start=pos,
                end=pos + len(highlighted_range.get_text()),
            )
            comments.append(
                comment.get_dict(self.config.include_author, self.config.include_date)
            )
        return comments

    def extract_comments_from_docx(self, docx_path: str) -> tuple[List[Dict], Section, Section]:
        """Extract comments directly from the Word document's XML structure."""
        # Read XML files
        comments_root, comments_extend_root, doc_root = self._read_docx_file(docx_path)

        # Extract comment ranges and their text
        comment_id_to_range, full_text = self._extract_highlight_ranges(doc_root)

        fb_section = self.extract_text_between_tokens(full_text, self.config.fb_start_token, self.config.fb_end_token)
        section_start = fb_section.start

        prompt_section = self.extract_text_between_tokens(full_text, self.config.prompt_start_token, self.config.prompt_end_token)
        prompt_start = prompt_section.start

        # Get sub-comments (replies)
        sub_comments = self._get_sub_comments(comments_extend_root)

        comments = []
        if comments_root is None:
            logger.warning("Skip document. No comments found")
        else:
            # Extract comments
            comments = self._extract_comments(
                comments_root, comment_id_to_range, section_start, sub_comments
            )

        return comments, fb_section, prompt_section
    
    def extract_table(self, input_file: str) -> List[Dict]:
        """Extract table from the document."""
        name_mapping = {
            "Item": "item",
            "Evaluation": "evaluation",
            "Optional comments": "optional_comments",
        }
        data = []

        document = Document(input_file)
        if len(document.tables) == 0:
            logger.error("No table found in the document")
            return data
        table = document.tables[0]
        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text.strip() for cell in row.cells)

            if i == 0:
                keys = tuple(name_mapping.get(cell.text.strip(), cell.text.strip()) for cell in row.cells)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)
        return data

    def process_document(self, file_path: str) -> dict[str, List[Dict]]:
        """Process a single document and return the JSON structure."""
        result = {ESSAY_PROMPT_KEY: None, ESSAY_TEXT_KEY: None, COMMENTS_KEY: []}
        try:
            # Process comments
            comments, fb_section, prompt_section = self.extract_comments_from_docx(file_path)
            general_feedbacks = self.extract_table(file_path)
            prompt_text = prompt_section.stripped_text if prompt_section else ""
            text = fb_section.stripped_text if fb_section else ""
            if not text:
                logger.warning("Could not find tokens in %s", file_path)
                return result

            result[ESSAY_PROMPT_KEY] = prompt_text
            result[ESSAY_TEXT_KEY] = text
            result[COMMENTS_KEY] = comments
            result[GENERAL_FEEDBACK_KEY] = general_feedbacks
            return result

        except Exception as e:
            logger.error("Error processing %s: %s", file_path, str(e))
            return result
