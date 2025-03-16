import json
import logging
from docx.api import Document

logger = logging.getLogger(__name__)

input_file = "data/input/samples/sample 1.docx"


def extract_table(input_file: str):
    data = []

    document = Document(input_file)
    if len(document.tables) == 0:
        logger.error("No table found in the document")
        return data
    table = document.tables[0]
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text.strip() for cell in row.cells)

        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return data

def main():
    # Example usage with tokens
    start_token = "Feedback on topic responsiveness:"  # Replace with your actual tokens
    end_token = None
    table_data = extract_table(input_file)
    print(json.dumps(table_data, indent=2))


if __name__ == "__main__":
    main()